import streamlit as st
import requests
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
import re
import csv
import io
from docx import Document
from docx.shared import RGBColor

# ==========================================
# 1. WEB SCRAPING FUNCTION
# ==========================================
def scrape_teacher_data(url):
    """
    Scrapes teacher data from the given URL.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.exceptions.RequestException:
        return []

    soup = BeautifulSoup(response.content, 'html.parser')
    teachers_data = []

    # STEP 1: Find all profile links from the main grid
    base_url = "https://ku.ac.bd"
    profile_links = []
    
    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        if '/faculty/' in href and not any(exclude in href.lower() for exclude in ['inservice', 'former', 'onleave']):
            full_url = href if href.startswith('http') else base_url + href
            if full_url not in profile_links:
                profile_links.append(full_url)

    # Automatically grab the discipline short code as a fallback
    department_short = url.split('/')[-2].upper()
    
    # Try to extract the full discipline name from the webpage title
    discipline_full = department_short
    title_tag = soup.find('title')
    if title_tag:
        title_text = title_tag.text
        for sep in ['|', '-']:
            if sep in title_text:
                parts = title_text.split(sep)
                for part in parts:
                    if "Discipline" in part or "School" in part or "Institute" in part:
                        discipline_full = part.strip()
                        break

    # STEP 2: Visit each profile to get details
    progress_text = "Scraping profiles. Please wait..."
    my_bar = st.progress(0, text=progress_text)
    
    for i, profile_url in enumerate(profile_links, 1):
        try:
            # Update web progress bar
            progress = int((i / len(profile_links)) * 100)
            my_bar.progress(progress, text=f"Scraping profile {i} of {len(profile_links)}...")
            
            prof_response = requests.get(profile_url, headers=headers)
            prof_soup = BeautifulSoup(prof_response.content, 'html.parser')
            
            name = "N/A"
            title_tag = prof_soup.find('title')
            if title_tag:
                name = title_tag.text.split('-')[0].strip()

            designation = "N/A"
            roles = ['Professor', 'Associate Professor', 'Assistant Professor', 'Lecturer']
            for tag in prof_soup.find_all(['p', 'span', 'h4', 'h5', 'div', 'li']):
                text = tag.get_text(strip=True)
                if any(role in text for role in roles):
                    for role in roles:
                        if role in text:
                            designation = role
                            break
                if designation != "N/A":
                    break

            email = "N/A"
            for a_tag in prof_soup.find_all('a', href=True):
                if 'mailto:' in a_tag['href']:
                    email = a_tag['href'].replace('mailto:', '').strip()
                    break
            
            if email == "N/A":
                email_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', prof_soup.get_text())
                if email_match:
                    email = email_match.group(0)

            teachers_data.append([name, designation, discipline_full, email])
            
        except Exception:
            continue
            
    my_bar.empty() # Clear progress bar when done
    return teachers_data

# ==========================================
# 2. PDF GENERATION FUNCTION (In-Memory)
# ==========================================
def create_id_cards_pdf(data):
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = styles['Heading1']
    title_style.alignment = 1 
    elements.append(Paragraph("Faculty ID Cards", title_style))
    elements.append(Spacer(1, 20))
    
    cards_data = []
    row = []
    
    for item in data:
        name, designation, discipline, email = item
        
        card_content = [
            [Paragraph(f" {name}", styles['Center'])],
            [Paragraph(f" {designation}", styles['Center '])],
            [Paragraph(f" {discipline}", styles['Center '])],
            [Spacer(1, 15)],
            [Paragraph("<b>Khulna University</b>", styles['Center'])]
        ]
        
        card_table = Table(card_content, colWidths=[230])
        card_table.setStyle(TableStyle([
            ('BOX', (0,0), (-1,-1), 2, colors.HexColor('#34495e')), 
            ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor("#ffffff")), 
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 12),
            ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ]))
        
        row.append(card_table)
        if len(row) == 2:
            cards_data.append(row)
            row = []
            
    if row:
        row.append("")
        cards_data.append(row)

    if cards_data:
        grid_table = Table(cards_data, colWidths=[260, 260], rowHeights=[140]*len(cards_data))
        grid_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        elements.append(grid_table)

    pdf.build(elements)
    return buffer.getvalue()

# ==========================================
# 3. CSV GENERATION FUNCTION (In-Memory)
# ==========================================
def create_csv(data):
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['Name', 'Designation', 'Discipline', 'Email'])
    writer.writerows(data)
    return buffer.getvalue().encode('utf-8')

# ==========================================
# 4. WORD DOCX FUNCTION (In-Memory)
# ==========================================
def create_docx(data):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading('Teacher Data', 0)

    for item in data:
        name, designation, discipline, email = item
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        p = cell.paragraphs[0]
        p.add_run(f"{name}\n")
        p.add_run(f"{designation}\n")
        p.add_run(f"{discipline},\n\n")
        
        uni_run = p.add_run("Khulna University")
        uni_run.bold = True        
        doc.add_paragraph() 

    doc.save(buffer)
    return buffer.getvalue()

# ==========================================
# 5. STREAMLIT WEB APP UI
# ==========================================
st.set_page_config(page_title="KU Faculty Scraper", page_icon="🎓")

st.title("🎓 KU Faculty Scraper")
st.write("Extract faculty data from Khulna University website and export it as ID cards or spreadsheets.")

# User Inputs
discipline_code = st.text_input("Enter Discipline Short Name:", value="fwt", help="Example: 'fwt' for Forestry and Wood Technology").strip().lower()
format_choice = st.selectbox("Select Export Format:", ["PDF", "Word Doc", "CSV (Spreadsheet)"])

if st.button("Scrape Data & Generate File"):
    if not discipline_code:
        st.warning("Please enter a discipline code.")
    else:
        TARGET_URL = f"https://ku.ac.bd/discipline/{discipline_code}/faculties"
        
        with st.spinner("Connecting to Khulna University website..."):
            scraped_data = scrape_teacher_data(TARGET_URL)

        if not scraped_data:
            st.error("Failed to scrape data. Please check if the discipline short name is correct.")
        else:
            st.success(f"Successfully grabbed {len(scraped_data)} faculty profiles!")
            
            # Generate the requested file in memory
            if format_choice == "PDF (ID Cards)":
                file_bytes = create_id_cards_pdf(scraped_data)
                mime_type = "application/pdf"
                file_ext = "pdf"
            elif format_choice == "Word Doc (ID Cards)":
                file_bytes = create_docx(scraped_data)
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_ext = "docx"
            else:
                file_bytes = create_csv(scraped_data)
                mime_type = "text/csv"
                file_ext = "csv"
            
            # Provide the download button
            filename = f"KU_{discipline_code.upper()}_Faculty.{file_ext}"
            st.download_button(
                label=f"⬇️ Download {filename}",
                data=file_bytes,
                file_name=filename,
                mime=mime_type
            )